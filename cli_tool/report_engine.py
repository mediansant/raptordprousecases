"""
RaptorDB Pro Readiness Analyzer — Built-in Report Engine

Generates a structured Markdown readiness report from collected data.
Rule-based scoring only — no LLM required.

Composite score per table (max 40 pts):
  Row count tier     : >10M=10, >1M=7, >100K=4, else 1
  Report density tier: >20=10, >10=7, >5=4, else 1
  PA indicator tier  : >5=10, >2=7, >0=4, else 0
  Slow tx hit tier   : >50=10, >10=7, >0=4, else 0
"""

import re
import tempfile
from datetime import datetime
from typing import Dict, List

import pandas as pd


# =============================================================================
# Scoring helpers
# =============================================================================

def _score_row_count(rows: int) -> int:
    if rows > 10_000_000: return 10
    if rows > 1_000_000:  return 7
    if rows > 100_000:    return 4
    return 1


def _score_report_density(reports: int) -> int:
    if reports > 20: return 10
    if reports > 10: return 7
    if reports > 5:  return 4
    return 1


def _score_pa_indicators(pa: int) -> int:
    if pa > 5: return 10
    if pa > 2: return 7
    if pa > 0: return 4
    return 0


def _score_slow_tx(hits: int) -> int:
    if hits > 50: return 10
    if hits > 10: return 7
    if hits > 0:  return 4
    return 0


# =============================================================================
# Internal helpers
# =============================================================================

def _get_props(results: Dict) -> Dict[str, str]:
    """Return system properties as a plain {name: value} dict."""
    df = results.get("system_properties")
    if df is None or df.empty or "name" not in df.columns:
        return {}
    return dict(zip(df["name"].astype(str), df["value"].astype(str)))


def _get_slow_hits_by_table(results: Dict) -> Dict[str, int]:
    """
    Parse slow_transaction_summary URLs to infer table names.
    Returns {table_name: total_hit_count}.
    """
    slow_df = results.get("slow_transaction_summary")
    hits: Dict[str, int] = {}
    if slow_df is None or slow_df.empty or "url" not in slow_df.columns:
        return hits
    count_col = (slow_df["count"] if "count" in slow_df.columns
                 else pd.Series([1] * len(slow_df)))
    for url, cnt in zip(slow_df["url"], count_col):
        if "/table/" in str(url):
            tbl = str(url).split("/table/")[-1].split("?")[0].split("/")[0]
            if tbl:
                hits[tbl] = hits.get(tbl, 0) + int(cnt)
    return hits


def _build_ranked_table(results: Dict) -> pd.DataFrame:
    """
    Score and rank every table that appears in at least one data dimension.
    Returns a DataFrame sorted descending by composite Score.
    """
    # Build dimension lookups
    row_counts: Dict[str, int] = {}
    rc_df = results.get("core_table_row_counts")
    if rc_df is not None and not rc_df.empty and "table_name" in rc_df.columns:
        row_counts = dict(zip(rc_df["table_name"], rc_df["row_count"]))

    report_counts: Dict[str, int] = {}
    rpt_df = results.get("report_table_summary")
    if rpt_df is not None and not rpt_df.empty and "table" in rpt_df.columns:
        report_counts = dict(zip(rpt_df["table"], rpt_df["report_count"]))

    pa_counts: Dict[str, int] = {}
    pa_df = results.get("pa_table_summary")
    if pa_df is not None and not pa_df.empty and "table_name" in pa_df.columns:
        pa_counts = dict(zip(pa_df["table_name"], pa_df["indicator_count"]))

    slow_hits = _get_slow_hits_by_table(results)

    all_tables = set(row_counts) | set(report_counts) | set(pa_counts)
    rows_list = []
    for table in all_tables:
        rows    = row_counts.get(table, 0)
        reports = report_counts.get(table, 0)
        pa      = pa_counts.get(table, 0)
        slow    = slow_hits.get(table, 0)
        score   = (_score_row_count(rows) + _score_report_density(reports)
                   + _score_pa_indicators(pa) + _score_slow_tx(slow))

        reasons = []
        if rows > 10_000_000:
            reasons.append(f"{rows:,} rows — columnar indexing critical at this scale")
        elif rows > 0:
            reasons.append(f"{rows:,} rows")
        if reports > 10:
            reasons.append(f"{reports} reports drive frequent scans")
        elif reports > 0:
            reasons.append(f"{reports} reports reference this table")
        if pa > 5:
            reasons.append(f"{pa} PA indicators — heavy periodic aggregation")
        elif pa > 0:
            reasons.append(f"{pa} PA indicators")
        if slow > 0:
            reasons.append(f"{slow} slow-transaction URL hits")

        rows_list.append({
            "Table":           table,
            "Rows":            rows,
            "Reports":         reports,
            "PA Indicators":   pa,
            "Slow Hits":       slow,
            "Score":           score,
            "Why RaptorDB Pro": "; ".join(reasons) or "Baseline candidate",
        })

    if not rows_list:
        return pd.DataFrame()

    df = (pd.DataFrame(rows_list)
          .sort_values("Score", ascending=False)
          .reset_index(drop=True))
    df.index = df.index + 1
    df.index.name = "Rank"
    return df


def _fmt(n) -> str:
    """Format a number with commas, or return 'N/A'."""
    try:
        return f"{int(n):,}"
    except (ValueError, TypeError):
        return "N/A"


def _md_table(df: pd.DataFrame) -> str:
    """Render a DataFrame as a Markdown pipe-table string."""
    if df.empty:
        return "_No data available._"
    cols   = list(df.columns)
    header = "| " + " | ".join(str(c) for c in cols) + " |"
    sep    = "| " + " | ".join("---" for _ in cols) + " |"
    body   = []
    for _, row in df.iterrows():
        cells = " | ".join("" if pd.isna(v) else str(v) for v in row)
        body.append(f"| {cells} |")
    return "\n".join([header, sep] + body)


# =============================================================================
# Health score
# =============================================================================

def _overall_health(issues: List[Dict]) -> tuple:
    """Return (label, emoji, description) based on highest severity present."""
    severities = {i["severity"] for i in issues}
    if "CRITICAL" in severities:
        return "Red",   "🔴", "Critical performance or configuration issues detected"
    if "WARNING" in severities:
        return "Amber", "🟡", "Warnings present — performance improvements available"
    return "Green", "🟢", "No critical issues — healthy baseline for RaptorDB Pro benchmarking"


# =============================================================================
# Section generators
# =============================================================================

def _section_executive_summary(results: Dict, issues: List[Dict]) -> str:
    props  = _get_props(results)
    rc_df  = results.get("core_table_row_counts")
    inv_df = results.get("table_inventory")
    health_label, health_emoji, health_desc = _overall_health(issues)

    build   = props.get("glide.buildtag", props.get("glide.buildname", "Unknown"))
    db_type = props.get("glide.db.type", "") or props.get("glide.db.rdbms", "Unknown")
    rdb_std = props.get("glide.raptordb.enabled", "false")
    rdb_pro = props.get("glide.raptordb.pro.enabled", "false")
    wdf     = props.get("sn_data_fabric.enabled", "false")
    pa_on   = props.get("glide.pa.enabled", "false")

    total_tables  = len(inv_df) if inv_df is not None and not inv_df.empty else "Unknown"
    top_tables_md = ""
    if rc_df is not None and not rc_df.empty:
        top_tables_md = _md_table(
            rc_df.head(5)[["table_name", "row_count"]]
            .rename(columns={"table_name": "Table", "row_count": "Row Count"})
            .assign(**{"Row Count": lambda d: d["Row Count"].apply(_fmt)})
            .reset_index(drop=True)
        )

    critical_cnt = sum(1 for i in issues if i["severity"] == "CRITICAL")
    warning_cnt  = sum(1 for i in issues if i["severity"] == "WARNING")
    info_cnt     = sum(1 for i in issues if i["severity"] == "INFO")

    if str(rdb_pro).lower() == "true":
        rdb_status = "RaptorDB Pro already enabled — focus on optimization."
    elif str(rdb_std).lower() == "true":
        rdb_status = "Standard RaptorDB active — upgrade to Pro for columnar HTAP."
    else:
        rdb_status = "Neither RaptorDB standard nor Pro is enabled — clean migration target."

    return "\n".join([
        "## 1. Executive Summary",
        "",
        "### Instance Profile",
        "",
        "| Property | Value |",
        "|---|---|",
        f"| Build | `{build}` |",
        f"| Database Engine | `{db_type}` |",
        f"| Total Tables in Inventory | {_fmt(total_tables)} |",
        f"| RaptorDB (standard) enabled | `{rdb_std}` |",
        f"| RaptorDB Pro enabled | `{rdb_pro}` |",
        f"| Workflow Data Fabric | `{wdf}` |",
        f"| Performance Analytics | `{pa_on}` |",
        "",
        "### Largest Tables",
        "",
        top_tables_md or "_Row count data not collected._",
        "",
        "### Overall Health",
        "",
        f"**{health_emoji} {health_label}** — {health_desc}",
        "",
        "| Severity | Count |",
        "|---|---|",
        f"| 🔴 CRITICAL | {critical_cnt} |",
        f"| 🟡 WARNING  | {warning_cnt} |",
        f"| 🔵 INFO     | {info_cnt} |",
        "",
        f"**RaptorDB Status:** {rdb_status}",
    ])


def _section_top10_tables(results: Dict) -> str:
    ranked = _build_ranked_table(results)
    if ranked.empty:
        return "## 2. Top 10 RaptorDB Pro Target Tables\n\n_Insufficient data._\n"

    top10 = ranked.head(10).copy()
    top10["Rows"]  = top10["Rows"].apply(_fmt)
    top10["Score"] = top10["Score"].apply(lambda s: f"{s} / 40")

    return "\n".join([
        "## 2. Top 10 RaptorDB Pro Target Tables",
        "",
        ("Tables ranked by composite score: "
         "**row-count tier** + **report-density tier** + **PA-indicator tier** + **slow-tx tier** (max 40 pts)."),
        "",
        _md_table(top10.reset_index()),
        "",
        "> **Score guide:** 30–40 = critical candidate · 20–29 = strong · 10–19 = moderate · <10 = baseline",
    ])


def _section_dashboard_hotspots(results: Dict) -> str:
    rpt_df = results.get("report_table_summary")
    pa_df  = results.get("pa_table_summary")
    rc_df  = results.get("core_table_row_counts")

    row_counts: Dict[str, int] = {}
    if rc_df is not None and not rc_df.empty and "table_name" in rc_df.columns:
        row_counts = dict(zip(rc_df["table_name"], rc_df["row_count"]))

    # Cross-reference: table → {reports, pa, row_count}
    tables: Dict[str, Dict] = {}
    if rpt_df is not None and not rpt_df.empty and "table" in rpt_df.columns:
        for _, row in rpt_df.iterrows():
            t = row["table"]
            tables.setdefault(t, {"reports": 0, "pa": 0})
            tables[t]["reports"] = int(row["report_count"])

    if pa_df is not None and not pa_df.empty and "table_name" in pa_df.columns:
        for _, row in pa_df.iterrows():
            t = row["table_name"]
            tables.setdefault(t, {"reports": 0, "pa": 0})
            tables[t]["pa"] = int(row["indicator_count"])

    if not tables:
        return "## 3. Dashboard & Report Hotspots\n\n_No report or PA data collected._\n"

    rows_list = []
    for t, d in tables.items():
        rc = row_counts.get(t, 0)
        impact = ("HIGH"   if (d["reports"] > 10 or d["pa"] > 3) and rc > 500_000
                  else "MEDIUM" if (d["reports"] > 5 or d["pa"] > 1)
                  else "LOW")
        rows_list.append({
            "Table":         t,
            "Reports":       d["reports"],
            "PA Indicators": d["pa"],
            "Row Count":     _fmt(rc),
            "Query Impact":  impact,
        })

    hotspot_df = (
        pd.DataFrame(rows_list)
        .sort_values(["Query Impact", "Reports", "PA Indicators"],
                     ascending=[True, False, False])
        .head(15)
        .reset_index(drop=True)
    )

    return "\n".join([
        "## 3. Dashboard & Report Hotspots",
        "",
        ("Tables hit most frequently by reports and PA indicator collections. "
         "Tables with both high report counts and large row sizes are the best demo scenarios "
         "for RaptorDB Pro's columnar query engine."),
        "",
        _md_table(hotspot_df),
        "",
        ("**Demo insight:** `HIGH` impact tables have >10 reports _and_ >500 K rows. "
         "A columnar index on the most-filtered field reduces dashboard load times by **5–20×** "
         "for aggregation queries."),
    ])


def _section_slow_transactions(results: Dict) -> str:
    slow_df = results.get("slow_transaction_summary")
    lines   = ["## 4. Slow Transaction Analysis", ""]

    if slow_df is None or slow_df.empty:
        lines += [
            "_No slow transactions recorded in the last 7 days._",
            "",
            ("This may mean the instance performs well, the slow-query threshold is set high, "
             "or `syslog_transaction` logging is disabled. "
             "Check `glide.db.slow_query_threshold`."),
        ]
        return "\n".join(lines)

    display = slow_df.copy().head(20)

    # Infer table name from URL pattern
    if "url" in display.columns:
        display["Inferred Table"] = display["url"].apply(
            lambda u: (str(u).split("/table/")[-1].split("?")[0].split("/")[0]
                       if "/table/" in str(u) else "—")
        )

    # Impact score = count × avg_response_ms
    if "count" in display.columns and "avg_response_ms" in display.columns:
        display["Impact Score"] = (
            pd.to_numeric(display["count"],           errors="coerce").fillna(0)
            * pd.to_numeric(display["avg_response_ms"], errors="coerce").fillna(0)
        ).astype(int)
        display = display.sort_values("Impact Score", ascending=False)

    show_cols = [c for c in
                 ["url", "Inferred Table", "count", "avg_response_ms",
                  "max_response_ms", "p90_response_ms", "Impact Score"]
                 if c in display.columns]

    lines += [
        (f"**{len(slow_df)} slow URL patterns** captured in the last 7 days (threshold >5 s). "
         "Ranked by `count × avg_response_ms` (Impact Score)."),
        "",
        _md_table(display[show_cols].reset_index(drop=True)),
        "",
        ("**RaptorDB Pro opportunity:** Each slow URL is a benchmarking candidate. "
         "Capture the exact query behind the top 3 patterns, then replay on a "
         "RaptorDB Pro-enabled instance for a direct before/after comparison."),
    ]
    return "\n".join(lines)


def _section_cmdb(results: Dict) -> str:
    cmdb_summary = results.get("cmdb_summary")
    ci_classes   = results.get("cmdb_ci_classes")
    lines        = ["## 5. CMDB Assessment", ""]

    if cmdb_summary is None or cmdb_summary.empty:
        lines.append("_CMDB summary not collected._")
        return "\n".join(lines)

    row        = cmdb_summary.iloc[0]
    total_cis  = int(row.get("total_cis", 0))
    ci_count   = int(row.get("ci_classes", 0))
    total_rels = int(row.get("total_relationships", 0))

    if total_cis > 500_000:
        maturity = "Enterprise CMDB — complex graph traversals are frequent"
    elif total_cis > 100_000:
        maturity = "Mid-size CMDB — growing relationship complexity"
    else:
        maturity = "Smaller CMDB — relationship queries manageable today"

    lines += [
        "| Metric | Value |",
        "|---|---|",
        f"| Total CIs | {_fmt(total_cis)} |",
        f"| CI Classes | {ci_count} |",
        f"| Total Relationships | {_fmt(total_rels)} |",
        f"| CMDB Maturity | {maturity} |",
        "",
    ]

    if ci_classes is not None and not ci_classes.empty:
        top10 = ci_classes.head(10).copy()
        if "count" in top10.columns:
            top10["count"] = top10["count"].apply(_fmt)
        lines += [
            "### Top 10 CI Classes by Count",
            "",
            _md_table(top10.reset_index(drop=True)),
            "",
        ]

    lines += [
        "### RaptorDB Pro Relevance",
        "",
        ("CMDB graph queries involve multi-table joins across `cmdb_ci`, `cmdb_rel_ci`, "
         "and dozens of extension tables. RaptorDB Pro's columnar engine accelerates:"),
        "- **Impact analysis** — traverse N hops of CI relationships",
        "- **Service dependency maps** — aggregate all CIs under a service tree",
        "- **Compliance dashboards** — filter/aggregate across CI classes and states",
        "",
        (f"With {_fmt(total_rels)} relationships, each dependency traversal is a candidate "
         f"for **5–30×** improvement with columnar indexing on `cmdb_rel_ci`."),
    ]
    return "\n".join(lines)


def _section_indexes(results: Dict) -> str:
    idx_df  = results.get("indexed_fields")
    comp_df = results.get("composite_indexes")
    lines   = ["## 6. Index Optimization Opportunities", ""]

    if idx_df is not None and not idx_df.empty and "name" in idx_df.columns:
        per_table    = (idx_df.groupby("name").size()
                        .reset_index(name="index_count")
                        .sort_values("index_count", ascending=False))
        over_indexed = per_table[per_table["index_count"] > 15]

        lines += [
            f"**Total indexed fields sampled:** {len(idx_df)}",
            "",
            "### Over-Indexed Tables (>15 indexed fields)",
            "",
        ]
        if not over_indexed.empty:
            lines.append(_md_table(over_indexed.head(10).reset_index(drop=True)))
            lines += [
                "",
                (f"**{len(over_indexed)} tables** carry >15 secondary indexes. "
                 "Each index adds write amplification on INSERT/UPDATE. "
                 "RaptorDB Pro's columnar engine can replace multiple secondary indexes "
                 "with a single columnar scan — reducing write overhead by **10–30%** per removed index."),
            ]
        else:
            lines.append("_No tables with >15 indexed fields found in the sampled data._")
        lines.append("")

    comp_count = len(comp_df) if comp_df is not None and not comp_df.empty else 0
    lines += [
        "### Composite Indexes",
        "",
        f"**{comp_count}** composite indexes found instance-wide.",
        "",
        ("Composite indexes were built to accelerate specific query patterns on MariaDB. "
         "With RaptorDB Pro's columnar engine most can be dropped or consolidated:"),
        "- Columnar storage makes range scans and aggregations fast without explicit indexes",
        "- Each dropped write-amplifying index reduces INSERT/UPDATE overhead",
        "- Estimate: **10–30%** write throughput improvement per removed composite index on high-write tables",
    ]
    if comp_df is not None and not comp_df.empty:
        lines += [
            "",
            "**Sample composite indexes:**",
            "",
            _md_table(comp_df.head(10).reset_index(drop=True)),
        ]
    return "\n".join(lines)


def _section_workload_profile(results: Dict) -> str:
    rpt_df  = results.get("report_table_summary")
    pa_df   = results.get("pa_indicators")
    jobs_df = results.get("scheduled_jobs")
    slow_df = results.get("slow_transaction_summary")

    total_reports = (int(rpt_df["report_count"].sum())
                     if rpt_df is not None and not rpt_df.empty
                        and "report_count" in rpt_df.columns else 0)
    total_pa      = len(pa_df)   if pa_df   is not None and not pa_df.empty   else 0
    total_jobs    = len(jobs_df) if jobs_df is not None and not jobs_df.empty else 0
    total_slow    = len(slow_df) if slow_df is not None and not slow_df.empty else 0

    read_proxy  = total_reports + total_pa
    write_proxy = max(total_jobs, 1)   # avoid div-by-zero
    total_proxy = read_proxy + write_proxy

    read_pct  = round(read_proxy  / total_proxy * 100)
    write_pct = round(write_proxy / total_proxy * 100)

    # HTAP suitability: higher read% + slow-query pressure = better fit
    htap_score = min(100, int(read_pct * 0.6 + min(40, total_slow * 0.4)))

    if htap_score >= 70:
        htap_label = "Excellent HTAP candidate — predominantly analytical workload"
    elif htap_score >= 40:
        htap_label = "Good HTAP candidate — mixed OLTP/OLAP"
    else:
        htap_label = "Primarily transactional — RaptorDB Pro helps most with reporting queries"

    return "\n".join([
        "## 7. Workload Profile",
        "",
        "### Read vs Write Estimate",
        "",
        "| Dimension | Count | Role |",
        "|---|---|---|",
        f"| Reports (read-load proxy) | {total_reports:,} | Analytical read |",
        f"| PA Indicators (periodic aggregations) | {total_pa:,} | OLAP read |",
        f"| Scheduled Jobs (background write proxy) | {total_jobs:,} | Write/mixed |",
        f"| Slow Transactions (last 7 days) | {total_slow:,} | Bottleneck signal |",
        "",
        f"**Estimated workload split:** ~{read_pct}% read-oriented / ~{write_pct}% write-oriented",
        "",
        "### HTAP Suitability",
        "",
        f"**HTAP Score: {htap_score} / 100** — {htap_label}",
        "",
        ("RaptorDB Pro's Hybrid Transactional/Analytical Processing architecture "
         "maintains separate columnar and row-store replicas so heavy analytical queries "
         "no longer contend with OLTP writes. "
         + ("This instance is an ideal HTAP deployment target."
            if htap_score >= 70
            else "This instance will see meaningful OLTP/reporting isolation benefits.")),
    ])


def _section_demo_scenarios(results: Dict, issues: List[Dict]) -> str:
    ranked  = _build_ranked_table(results)
    slow_df = results.get("slow_transaction_summary")
    pa_dash = results.get("pa_dashboards")
    cmdb_s  = results.get("cmdb_summary")
    scenarios = []

    # --- Scenarios 1–3: top-ranked tables ---
    for i, (_, row) in enumerate(ranked.head(3).iterrows(), 1):
        tbl     = row["Table"]
        rows    = _fmt(row["Rows"])
        reports = int(row["Reports"])
        pa      = int(row["PA Indicators"])

        if pa > 0:
            query_type  = "PA aggregation / dashboard load"
            test        = (f"Open a PA dashboard sourced from `{tbl}` "
                           f"({rows} rows, {pa} PA indicators). "
                           "Capture end-to-end dashboard load time.")
            expectation = "5–20× load time reduction with columnar snapshot queries"
        elif reports > 5:
            query_type  = "Multi-filter list report"
            test        = (f"Run the most complex report on `{tbl}` ({rows} rows, "
                           f"{reports} reports defined). "
                           "Record total query time in `syslog_transaction`.")
            expectation = "3–10× query time reduction with columnar index scan"
        else:
            query_type  = "Large table aggregation scan"
            test        = (f"Run a COUNT/GROUP BY query on `{tbl}` ({rows} rows). "
                           "Measure execution time via `sys_db_slow_query`.")
            expectation = "5–50× improvement on full-table aggregation"

        scenarios.append(
            f"### Demo {i}: `{tbl}` — {query_type}\n\n"
            f"**Setup:** {test}\n\n"
            f"**Expected improvement:** {expectation}\n\n"
            f"**Why this table:** {row['Why RaptorDB Pro']}"
        )

    # --- Scenario 4: slow transaction replay ---
    if slow_df is not None and not slow_df.empty and "url" in slow_df.columns:
        top   = slow_df.iloc[0]
        url   = top.get("url", "N/A")
        avg   = top.get("avg_response_ms", "N/A")
        cnt   = top.get("count", "N/A")
        scenarios.append(
            f"### Demo 4: Slow Transaction Replay — `{url}`\n\n"
            f"**Setup:** This URL pattern generated {cnt} slow transactions averaging {avg} ms. "
            f"Identify the underlying table query from the transaction log, "
            f"then replay it on a RaptorDB Pro-enabled instance with the same dataset.\n\n"
            f"**Expected improvement:** Columnar scan on filtered queries typically yields "
            f"10–50× improvement; produces a direct, customer-visible before/after number.\n\n"
            f"**Why this scenario:** Highest-frequency slow pattern — most visible to end users."
        )
    else:
        scenarios.append(
            "### Demo 4: Scheduled Report Benchmark\n\n"
            "**Setup:** Identify the longest-running nightly report from `sys_report`. "
            "Compare execution time on MariaDB vs RaptorDB Pro with identical data.\n\n"
            "**Expected improvement:** 5–15× on aggregation-heavy report queries.\n\n"
            "**Why this scenario:** Scheduled reports delay start-of-day dashboard availability."
        )

    # --- Scenario 5: CMDB or PA dashboards ---
    if cmdb_s is not None and not cmdb_s.empty:
        total_cis  = int(cmdb_s.iloc[0].get("total_cis", 0))
        total_rels = int(cmdb_s.iloc[0].get("total_relationships", 0))
        if total_cis > 100_000:
            scenarios.append(
                f"### Demo 5: CMDB Impact Analysis Query\n\n"
                f"**Setup:** Execute a service impact analysis across {_fmt(total_cis)} CIs "
                f"and {_fmt(total_rels)} relationships "
                f"(e.g., 'find all CIs affected by outage of Service X'). "
                f"Measure multi-hop traversal time end-to-end.\n\n"
                f"**Expected improvement:** CMDB graph traversals show **10–30×** speedup "
                f"with columnar indexing on `cmdb_rel_ci`.\n\n"
                f"**Why this scenario:** CMDB queries are notoriously slow at enterprise scale — "
                f"a compelling, customer-visible win."
            )
        else:
            scenarios.append(_default_scenario_5())
    elif pa_dash is not None and not pa_dash.empty:
        scenarios.append(
            f"### Demo 5: PA Dashboard Parallel Load Test\n\n"
            f"**Setup:** Load {len(pa_dash)} PA dashboards simultaneously, "
            f"simulating start-of-day executive usage. Measure time to first render per dashboard.\n\n"
            f"**Expected improvement:** Concurrent columnar queries vs. serialised row-store scans "
            f"— typically **5–10×** reduction in total render time.\n\n"
            f"**Why this scenario:** Executive dashboard slowness is the most visible performance pain point."
        )
    else:
        scenarios.append(_default_scenario_5())

    lines = ["## 8. Recommended Demo Scenarios", ""]
    lines.extend(scenarios)
    return "\n".join(lines)


def _default_scenario_5() -> str:
    return (
        "### Demo 5: Audit Log Compliance Query\n\n"
        "**Setup:** Run a date-range audit query: "
        "'Show all changes to the Incident table in the last 90 days by user X'. "
        "Measure query time against `sys_audit` + `sys_journal_field`.\n\n"
        "**Expected improvement:** Audit tables are the fastest-growing in ServiceNow. "
        "Columnar range scans on date fields show **20–100×** speedup vs MariaDB row scans.\n\n"
        "**Why this scenario:** Compliance teams run these queries regularly — "
        "easy to demonstrate, with clear business value."
    )


def _section_next_steps(results: Dict, issues: List[Dict]) -> str:
    props       = _get_props(results)
    slow_thresh = props.get("glide.db.slow_query_threshold", "")
    pa_on       = props.get("glide.pa.enabled", "false")
    rdb_pro     = props.get("glide.raptordb.pro.enabled", "false")

    steps = []
    n = 1

    if str(rdb_pro).lower() == "true":
        steps.append(
            f"{n}. **RaptorDB Pro is already enabled** — proceed directly to benchmarking "
            f"the top-scored tables using the demo scenarios above."
        )
    else:
        steps.append(
            f"{n}. **Request RaptorDB Pro trial activation** from your ServiceNow account team, "
            f"targeting the top-scored tables identified in Section 2."
        )
    n += 1

    if not slow_thresh or slow_thresh == "0":
        steps.append(
            f"{n}. **Enable slow query logging**: set `glide.db.slow_query_threshold` "
            f"to `2000` (2 seconds) to capture more query patterns for benchmarking."
        )
    else:
        steps.append(
            f"{n}. **Expand slow query capture**: current threshold is `{slow_thresh} ms`. "
            f"Consider lowering to `1000 ms` to capture more patterns before the PoV."
        )
    n += 1

    if str(pa_on).lower() == "true":
        steps.append(
            f"{n}. **Identify top PA indicators for benchmarking**: from the PA table summary, "
            f"select the 3–5 indicators on the largest source tables as before/after test cases."
        )
    else:
        steps.append(
            f"{n}. **Enable Performance Analytics** and run an initial snapshot collection "
            f"to build a baseline before the PoV."
        )
    n += 1

    ranked = _build_ranked_table(results)
    if not ranked.empty:
        top_tbl = ranked.iloc[0]["Table"]
        steps.append(
            f"{n}. **Establish row-count baseline**: document current row counts for the top 10 "
            f"tables (especially `{top_tbl}`) before and after enabling RaptorDB Pro."
        )
        n += 1

    steps += [
        (f"{n}. **Capture execution plans**: use `sys_db_slow_query` and `syslog_transaction` "
         f"to record SQL execution plans for the top 5 slow queries before the PoV."),
        (f"{n+1}. **Schedule a PoV kick-off**: share this report with the ServiceNow RaptorDB Pro team. "
         f"Agree on 3 benchmark scenarios and success criteria (target improvement %)."),
        (f"{n+2}. **Quantify business impact**: for each demo scenario estimate daily user count "
         f"and query frequency to translate query speedup into hours saved per day."),
    ]

    return "\n".join([
        "## 9. Next Steps",
        "",
        "Concrete actions to prepare for a RaptorDB Pro Proof of Value:",
        "",
    ] + steps)


# =============================================================================
# Public API
# =============================================================================

def generate_report(results: Dict, issues: List[Dict]) -> str:
    """
    Generate the full RaptorDB Pro Readiness Report as a Markdown string.
    All scoring is rule-based; no LLM or external service is required.
    """
    now = datetime.now().strftime("%Y-%m-%d %H:%M")
    divider = "\n\n---\n\n"

    sections = [
        f"# RaptorDB Pro Readiness Report",
        f"",
        f"_Generated: {now}_",
        f"",
        "---",
        "",
        _section_executive_summary(results, issues),
        divider,
        _section_top10_tables(results),
        divider,
        _section_dashboard_hotspots(results),
        divider,
        _section_slow_transactions(results),
        divider,
        _section_cmdb(results),
        divider,
        _section_indexes(results),
        divider,
        _section_workload_profile(results),
        divider,
        _section_demo_scenarios(results, issues),
        divider,
        _section_next_steps(results, issues),
        "",
        "---",
        "",
        f"_Report generated by RaptorDB Pro Readiness Analyzer · {now}_",
    ]
    return "\n".join(sections)


def generate_claude_prompt(results: Dict, issues: List[Dict]) -> str:
    """
    Build a single optimized prompt for Claude to generate a rich natural-language
    readiness report. Large DataFrames are truncated to 50 rows; raw data is summarised
    rather than fully dumped to keep the prompt within token limits.
    """
    lines = [
        "You are a ServiceNow database performance expert and RaptorDB Pro specialist.",
        "Based on the following data collected from a live ServiceNow instance running on MariaDB,",
        "generate a comprehensive RaptorDB Pro Readiness Report in Markdown format.",
        "",
        "Do NOT be generic — reference actual table names, row counts, and query patterns",
        "from the data below. Every claim should cite specific numbers.",
        "",
        "Required sections:",
        "1. **Executive Summary** — instance profile (build, DB type, RaptorDB/WDF status),",
        "   overall health (🔴 Red / 🟡 Amber / 🟢 Green with justification)",
        "2. **Top 10 RaptorDB Pro Target Tables** — ranked by query load,",
        "   with one paragraph of justification per table",
        "3. **Dashboard & Report Hotspots** — which reports and PA dashboards hit the largest",
        "   tables; identify the highest-impact demo scenarios",
        "4. **Slow Transaction Analysis** — map URL patterns to tables, rank by",
        "   frequency × avg response time, suggest specific queries to benchmark",
        "5. **CMDB Assessment** — CI count, relationship complexity, expected improvement",
        "   for graph traversal queries",
        "6. **Index Optimization Opportunities** — over-indexed tables, composite index",
        "   consolidation potential, estimated write overhead reduction",
        "7. **Workload Profile** — read vs write ratio estimate, HTAP suitability score",
        "   (0–100), explanation of score",
        "8. **Top 5 Demo Scenarios** — specific before/after test cases with exact table",
        "   names, realistic improvement factors, and why each is compelling to a customer",
        "9. **Next Steps** — concrete, prioritised actions to prepare for a RaptorDB Pro PoV",
        "",
        "Use these RaptorDB Pro benchmark references for improvement estimates:",
        "- PA aggregation queries on >1M-row tables:   10–50× faster",
        "- Full table scans on >10M-row tables:        5–100× faster with columnar indexing",
        "- CMDB relationship traversals:               5–30× faster",
        "- Concurrent dashboard load (HTAP isolation): 5–10× improvement",
        "- Audit log date-range scans:                 20–100× faster",
        "",
        "=" * 72,
        "INSTANCE DATA",
        "=" * 72,
        "",
    ]

    # System properties
    props_df = results.get("system_properties")
    if props_df is not None and not props_df.empty:
        lines.append("### System Properties")
        for _, row in props_df.iterrows():
            lines.append(f"  {row.get('name', '')}: {row.get('value', '')}")
        lines.append("")

    # Flagged issues (all of them — they're already concise)
    if issues:
        lines.append("### Flagged Issues")
        for issue in issues:
            lines.append(f"  [{issue['severity']}] {issue['title']}: {issue['detail']}")
            if issue.get("raptordb_relevance"):
                lines.append(f"    → RaptorDB: {issue['raptordb_relevance']}")
        lines.append("")

    # Table row counts (top 50 by size)
    rc_df = results.get("core_table_row_counts")
    if rc_df is not None and not rc_df.empty:
        lines.append("### Table Row Counts (top 50 by size)")
        lines.append(rc_df.head(50).to_string(index=False))
        lines.append("")

    # Report table summary (top 50)
    rpt_df = results.get("report_table_summary")
    if rpt_df is not None and not rpt_df.empty:
        lines.append("### Report Count by Table (top 50)")
        lines.append(rpt_df.head(50).to_string(index=False))
        lines.append("")

    # PA table summary (top 50)
    pa_df = results.get("pa_table_summary")
    if pa_df is not None and not pa_df.empty:
        lines.append("### PA Indicators by Table (top 50)")
        lines.append(pa_df.head(50).to_string(index=False))
        lines.append("")

    # Slow transaction summary (top 50)
    slow_df = results.get("slow_transaction_summary")
    if slow_df is not None and not slow_df.empty:
        lines.append("### Slow Transaction Patterns (top 50 by count)")
        lines.append(slow_df.head(50).to_string(index=False))
        lines.append("")

    # CMDB summary
    cmdb_df = results.get("cmdb_summary")
    if cmdb_df is not None and not cmdb_df.empty:
        lines.append("### CMDB Summary")
        lines.append(cmdb_df.to_string(index=False))
        lines.append("")

    # CMDB CI class distribution (top 20)
    ci_df = results.get("cmdb_ci_classes")
    if ci_df is not None and not ci_df.empty:
        lines.append("### CMDB CI Class Distribution (top 20)")
        lines.append(ci_df.head(20).to_string(index=False))
        lines.append("")

    # Composite indexes (top 50)
    idx_df = results.get("composite_indexes")
    if idx_df is not None and not idx_df.empty:
        lines.append(f"### Composite Indexes ({len(idx_df)} total; top 50 shown)")
        lines.append(idx_df.head(50).to_string(index=False))
        lines.append("")

    # Indexed fields — summarised by table (top 50)
    if_df = results.get("indexed_fields")
    if if_df is not None and not if_df.empty and "name" in if_df.columns:
        per_tbl = (if_df.groupby("name").size()
                   .reset_index(name="index_count")
                   .sort_values("index_count", ascending=False)
                   .head(50))
        lines.append("### Indexed Fields per Table — summarised (top 50)")
        lines.append(per_tbl.to_string(index=False))
        lines.append("")

    # Scheduled jobs (summarised by run_type)
    jobs_df = results.get("scheduled_jobs")
    if jobs_df is not None and not jobs_df.empty:
        lines.append(f"### Scheduled Jobs: {len(jobs_df)} active")
        if "run_type" in jobs_df.columns:
            by_type = jobs_df.groupby("run_type").size().reset_index(name="count")
            lines.append(by_type.to_string(index=False))
        lines.append("")

    # Table rotation
    rot_df = results.get("table_rotation")
    if rot_df is not None and not rot_df.empty:
        lines.append(f"### Table Rotation Rules: {len(rot_df)} active")
        lines.append(rot_df.to_string(index=False))
        lines.append("")

    lines += [
        "=" * 72,
        "END OF DATA",
        "=" * 72,
        "",
        ("Now generate the complete RaptorDB Pro Readiness Report following the structure above. "
         "Be specific, actionable, and professional. "
         "Target audience: ServiceNow platform admin + IT leadership."),
    ]

    return "\n".join(lines)


def export_docx_report(markdown_text: str, filepath: str) -> str:
    """
    Convert a Markdown report string to a .docx file using python-docx.
    Handles H1–H4 headings, bullet/numbered lists, pipe tables, block quotes,
    horizontal rules, and inline **bold** / `code` spans.
    Returns the filepath on success.
    """
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.oxml import OxmlElement

    doc = Document()
    doc.core_properties.title  = "RaptorDB Pro Readiness Report"
    doc.core_properties.author = "RaptorDB Pro Readiness Analyzer"

    # Readable margins
    for sec in doc.sections:
        sec.top_margin    = Inches(1.0)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin   = Inches(1.0)
        sec.right_margin  = Inches(1.0)

    def _rich_para(text: str, style: str = "Normal"):
        """Add a paragraph with inline **bold** and `code` support."""
        para   = doc.add_paragraph(style=style)
        tokens = re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text)
        for tok in tokens:
            if tok.startswith("**") and tok.endswith("**"):
                run = para.add_run(tok[2:-2])
                run.bold = True
            elif tok.startswith("`") and tok.endswith("`"):
                run = para.add_run(tok[1:-1])
                run.font.name = "Courier New"
                run.font.size = Pt(9)
            else:
                para.add_run(tok)
        return para

    def _add_table(header_line: str, data_lines: list):
        """Render a Markdown pipe-table as a Word table with a bold header row."""
        headers = [h.strip() for h in header_line.strip("|").split("|")]
        tbl     = doc.add_table(rows=1, cols=len(headers))
        tbl.style = "Table Grid"
        hdr_cells = tbl.rows[0].cells
        for j, h in enumerate(headers):
            hdr_cells[j].text = h
            for para in hdr_cells[j].paragraphs:
                for run in para.runs:
                    run.bold = True
        for data_line in data_lines:
            cells     = [c.strip() for c in data_line.strip("|").split("|")]
            row_cells = tbl.add_row().cells
            for j, c in enumerate(cells[: len(headers)]):
                row_cells[j].text = c
        doc.add_paragraph()  # spacing after table

    lines = markdown_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        # Headings
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)

        # Horizontal rule
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 90)

        # Markdown pipe-table: header row + separator row
        elif (line.startswith("|")
              and i + 1 < len(lines)
              and re.match(r"^\|[-| :]+\|", lines[i + 1])):
            i += 2  # skip separator
            data_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                data_lines.append(lines[i])
                i += 1
            _add_table(line, data_lines)
            continue  # i already advanced

        # Bullet list
        elif re.match(r"^[-*] ", line):
            _rich_para(line[2:], style="List Bullet")

        # Numbered list
        elif re.match(r"^\d+\. ", line):
            _rich_para(re.sub(r"^\d+\. ", "", line), style="List Number")

        # Block quote
        elif line.startswith("> "):
            para = doc.add_paragraph(style="Normal")
            run  = para.add_run(line[2:])
            run.italic = True
            para.paragraph_format.left_indent = Inches(0.35)

        # Non-empty paragraph
        elif line.strip():
            _rich_para(line)

        # Empty line → small spacer
        else:
            doc.add_paragraph()

        i += 1

    doc.save(filepath)
    return filepath
